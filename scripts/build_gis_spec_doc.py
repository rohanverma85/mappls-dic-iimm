from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "IIMM_GIS_KML_SHP_PRD_FRS_SRS_v0.1.docx"
NAVY = "104685"
NAVY_DARK = "0A2A50"
BLUE_LIGHT = "EAF1F8"
GOLD = "D99A00"
GREEN = "1F7A46"
GREY = "667085"
LINE = "D9DEE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_rule(paragraph, color=NAVY, size=16):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_note(doc, title, text, color=BLUE_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    p.add_run(title + "\n").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    hdr = table.rows[0]
    hdr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(font_size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        cells = row.cells
        for i, value in enumerate(values):
            set_cell_margins(cells[i])
            if row_index % 2:
                set_cell_shading(cells[i], "F7F9FC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    return table


def add_requirement_table(doc, requirements):
    return add_table(doc, ["ID", "Requirement", "Acceptance / evidence", "Priority"], requirements, [0.75, 3.25, 2.3, 0.65], 8.0)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = RGBColor.from_string("273142")
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.08
for name, size, color in (("Title", 30, NAVY_DARK), ("Heading 1", 20, NAVY_DARK), ("Heading 2", 14, NAVY), ("Heading 3", 11, GREY)):
    style = styles[name]
    style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.keep_with_next = True
styles["List Bullet"].font.name = "Aptos"
styles["List Bullet"].font.size = Pt(9.5)
styles["List Bullet 2"].font.name = "Aptos"
styles["List Bullet 2"].font.size = Pt(9)
styles["List Number"].font.name = "Aptos"
styles["List Number"].font.size = Pt(9.5)

header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hp.add_run("DIGITAL INDIA  /  IIMM PLATFORM")
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor.from_string(NAVY)
add_rule(hp, NAVY, 10)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Client review draft  •  v0.1  •  19 August 2026     |     ")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string(GREY)
page_field(fp)

# Cover / memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(26)
r = p.add_run("DIGITAL INDIA")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(NAVY)
r.font.letter_spacing = Pt(1.4)
p = doc.add_paragraph()
p.add_run("IIMM PLATFORM").bold = True
p.runs[0].font.size = Pt(12)
p.runs[0].font.color.rgb = RGBColor.from_string(GOLD)
title = doc.add_paragraph(style="Title")
title.add_run("GIS Asset & Network Import\nProduct, Functional and Software Requirements")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(10)
subtitle.add_run("KML / KMZ / Shapefile • Mappls operational map • geotagged defect lifecycle").bold = True
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
add_rule(subtitle, GOLD, 22)
doc.add_paragraph("A controlled, assumption-led extension to PRD_IIMM_Platform_v1.2 for client review and course correction. The implementation is isolated on a reversible feature branch and is not merged into the published production baseline.")
doc.add_paragraph()
add_table(doc, ["Document", "Value"], [
    ("Document ID", "IIMM-GIS-SPEC-001"),
    ("Version / status", "v0.1 • Client review draft"),
    ("Prepared", "19 August 2026 • Asia/Kolkata"),
    ("Implementation branch", "feature/kml-shp-gis-import"),
    ("Production baseline", "main • commit 2fa44ad"),
    ("Intended reviewers", "Client Product Owner, GIS Lead, Authority Operations, Security and Delivery teams"),
], [1.65, 5.35], 9)
doc.add_paragraph()
add_note(doc, "Decision status", "The three client asks are not specified in sufficient detail in PRD v1.2 or the supplied standalone prototype. This document deliberately converts stated intent into testable assumptions. Items marked DECISION REQUIRED must be confirmed before production merge.", "FFF3D6")
doc.add_page_break()

doc.add_heading("Document control", level=1)
add_table(doc, ["Version", "Date", "Author / owner", "Change", "Status"], [
    ("0.1", "19 Aug 2026", "IIMM delivery team", "Initial assumption-led PRD/FRS/SRS and implemented branch baseline", "For client review"),
    ("0.2", "TBD", "Client + delivery", "Decision updates and course correction", "Planned"),
    ("1.0", "TBD", "Client Product Owner", "Approved requirements baseline", "Planned"),
], [0.65, 0.85, 1.4, 3.45, 1.0])
doc.add_heading("Approval record", level=2)
add_table(doc, ["Role", "Name", "Decision", "Date / signature"], [
    ("Client Product Owner", "", "Approve / Approve with changes / Reject", ""),
    ("Client GIS Lead", "", "Approve / Approve with changes / Reject", ""),
    ("Client Security Lead", "", "Approve / Approve with changes / Reject", ""),
    ("Delivery Lead", "", "Accepted for implementation", ""),
], [1.55, 1.45, 2.75, 1.55])
doc.add_heading("How to review", level=2)
add_number(doc, "Review the source-gap statement and confirm that this extension is a new requirement rather than a correction to PRD v1.2.")
add_number(doc, "Resolve each DECISION REQUIRED item in section 10; record approved values in v0.2.")
add_number(doc, "Walk through the implemented feature branch using the acceptance scenarios in section 9.")
add_number(doc, "Approve merge, request changes on the same branch, or reject and retain the current production baseline.")
doc.add_heading("Contents", level=2)
for item in ["1. Executive summary", "2. Source traceability and gap analysis", "3. Assumptions and boundaries", "4. Product Requirements (PRD)", "5. Functional Requirements (FRS)", "6. Software Requirements (SRS)", "7. Security, privacy and audit", "8. Non-functional requirements", "9. Acceptance and traceability", "10. Client decisions required", "11. Rollback and change control"]:
    add_bullet(doc, item)
doc.add_page_break()

doc.add_heading("1. Executive summary", level=1)
doc.add_paragraph("The proposed extension makes spatial network data operational inside IIMM. An Authority User can import a KML, KMZ or zipped Shapefile, map source attributes, validate geometry on Mappls, and publish both a versioned GIS layer and tenant-scoped asset records. Makers and Checkers then see imported networks below defects, capture geotagged evidence, enforce project/asset proximity context, and complete the existing Maker–Checker defect-resolution chain.")
doc.add_heading("1.1 Client asks translated into outcomes", level=2)
add_table(doc, ["Client ask", "Proposed outcome", "Review status"], [
    ("Create assets using KML", "KML/KMZ/SHP upload → preview → field mapping → validated publish → one asset per feature", "Implemented on branch"),
    ("Add and resolve defect with geolocation, geotag and geofence", "Device coordinates + accuracy + media evidence; asset/project geofence context; Maker ATR; Checker verification", "Existing flow enhanced by imported asset context"),
    ("View defect on map over KML/SHP network", "Mappls vector basemap with tenant-scoped network layers, assets, geofences and selectable defects", "Implemented on branch"),
], [1.85, 4.2, 1.25])
doc.add_heading("1.2 Business benefits", level=2)
for text in [
    "Faster onboarding of existing road, building, utility and civic-asset inventories without manual re-entry.",
    "A single visual context for network ownership, defect proximity, field action and verification evidence.",
    "Controlled versioning and rollback so a bad GIS import does not silently corrupt the active network.",
    "Tenant, project and role boundaries preserved from the base IIMM architecture.",
    "A client-reviewable assumptions register that supports course correction before production merge.",
]: add_bullet(doc, text)
doc.add_heading("1.3 Recommended decision", level=2)
add_note(doc, "Recommendation", "Approve the branch for structured UAT, not immediate production merge. Confirm coordinate systems, attribute-mapping rules, import size/volume, geofence policy and authoritative-source ownership first.")
doc.add_page_break()

doc.add_heading("2. Source traceability and gap analysis", level=1)
doc.add_paragraph("The supplied PRD and standalone HTML were treated as authoritative for the prototype. They establish multi-tenancy, configurable asset types, geofenced Maker attendance, geo-tagged defect media, Citizen reporting, offline sync, Maker–Checker verification and Map/GIS-friendly architecture. They do not define GIS file ingestion or its operational rules.")
add_table(doc, ["Topic", "Source position", "Gap / consequence"], [
    ("Asset creation from KML/KMZ/SHP", "Not specified", "File formats, CRS, identity, mapping, validation, versioning and rollback require new requirements."),
    ("Defect geolocation / media", "Geo-tagged defect reporting and mobile evidence are in scope", "Accuracy thresholds, network snapping and out-of-geofence behaviour are not specified."),
    ("Geofencing", "Mobile-only Maker attendance is geofenced", "Defect-specific geofence enforcement is not defined; proposed as contextual validation."),
    ("Map view over underlying network", "Deep GIS/Bhuvan integration is out of prototype scope", "Local GeoJSON/KML/SHP overlays are an extension; no external government-system dependency is introduced."),
    ("Mappls credentials", "Not specified in PRD", "Current Web JS and reverse-geocode APIs use a Console static key; mGIS methods require OAuth credentials/token."),
], [1.55, 2.35, 3.4])
doc.add_heading("2.1 Traceability principle", level=2)
doc.add_paragraph("No proposed GIS rule is represented as an original PRD requirement. New requirements use the prefix GIS- and explicitly identify the assumption they implement. Existing defect, access, audit and offline requirements remain controlling where they intersect.")
doc.add_heading("2.2 Out of scope for this extension", level=2)
for text in [
    "Authoritative bidirectional sync with Mappls mGIS, Bhuvan, NIC, PFMS, GeM or third-party enterprise GIS.",
    "Editing source geometries in IIMM, topology repair, cartographic generalisation or survey-grade coordinate correction.",
    "Contract/agreement management, budget allocation and fund-availability checks.",
    "Native iOS/Android binaries; the responsive field experience remains a web prototype/PWA-ready implementation.",
    "Automatic rejection of genuine field defects solely because the device is outside a configured geofence.",
]: add_bullet(doc, text)
doc.add_page_break()

doc.add_heading("3. Assumptions and boundaries", level=1)
add_table(doc, ["ID", "Assumption proposed for v0.1", "Reason / course-correction path"], [
    ("ASM-01", "Accepted files: .kml, .kmz, and .zip containing a Shapefile set.", "Covers client request while keeping parsing client-side and provider-neutral."),
    ("ASM-02", "KML/KMZ coordinates are WGS84 (EPSG:4326). Shapefile ZIP must contain usable .prj projection metadata.", "Mappls and stored GeoJSON use longitude/latitude; client must confirm allowed source CRS list."),
    ("ASM-03", "Supported geometry: Point, LineString, Polygon and their multi-line/multi-polygon forms; GeometryCollection is rejected.", "Matches IIMM operational map and asset geometry model."),
    ("ASM-04", "Each valid feature creates or updates one asset instance and also remains in the published network layer.", "Provides searchable asset register plus visual network context."),
    ("ASM-05", "Source identity uses a selected field; otherwise asset_id, id, name, feature ID, then a deterministic geometry hash.", "Enables repeat import/upsert while avoiding silent duplicates."),
    ("ASM-06", "Authority Users publish and roll back imports. Makers, Checkers and Citizens can view only what their role and tenant allow.", "Preserves governance and separation of duties."),
    ("ASM-07", "KML ≤10 MB; KMZ/SHP ZIP ≤25 MB; maximum 5,000 valid features per transaction.", "Prototype-safe browser/server limits; scale test before changing."),
    ("ASM-08", "Publishing a replacement hides the previous layer, increments the version and retains a rollback link.", "Provides reversible network evolution."),
    ("ASM-09", "Default defect context: 75 m asset buffer; otherwise project radius. Outside results are flagged/audited, not silently blocked.", "Protects genuine incident capture and aligns with current prototype behaviour."),
    ("ASM-10", "Suggested network linkage/snapping tolerance is 25 m; original device coordinates remain authoritative evidence.", "Prevents coordinate mutation while offering operational context; not yet implemented in v0.1."),
    ("ASM-11", "Device GPS accuracy is recorded; policy threshold is configurable, initially 100 m warning and not hard rejection.", "Low-connectivity field conditions vary; client must set threshold."),
    ("ASM-12", "Mappls static key is allowed in Web JS and current reverse-geocode flow when both APIs are enabled. mGIS OAuth remains server-only and optional.", "Separates browser static credential from mGIS/private-layer access."),
], [0.65, 4.35, 2.35], 7.7)
doc.add_heading("3.1 Failure policy", level=2)
add_bullet(doc, "Parse and validate before publish; unsupported/empty geometry is shown as a validation note.")
add_bullet(doc, "Duplicate chosen source IDs block publishing until the user selects a unique field.")
add_bullet(doc, "Tenant/project/asset-type validation is repeated server-side; browser claims are never trusted alone.")
add_bullet(doc, "A published import is atomic at the IIMM transaction level. Partial publishing is not offered in v0.1.")
doc.add_page_break()

doc.add_heading("4. Product Requirements (PRD)", level=1)
doc.add_heading("4.1 Product goal", level=2)
doc.add_paragraph("Enable public-infrastructure authorities to convert existing GIS files into governed IIMM assets and network layers, then use that spatial context throughout defect capture, rectification and verification without weakening tenant isolation or the Maker–Checker model.")
doc.add_heading("4.2 Users and jobs", level=2)
add_table(doc, ["Persona", "Primary job", "Product entitlement"], [
    ("Authority User", "Onboard authoritative asset/network data and govern versions", "Upload, validate, map, publish, replace, roll back and audit"),
    ("External Maker", "Find assigned asset/defect and submit geotagged rectification evidence", "View network, navigate/focus, start work and submit ATR"),
    ("External Checker", "Validate issue and verify repair against location/network evidence", "View network, validate citizen defect and verify/reject ATR"),
    ("Citizen", "Report and track a nearby public-infrastructure issue", "Capture coordinates/media and see own issue status; no GIS import access"),
    ("Tenant Administrator", "Configure tenant and support securely", "No implicit access to tenant operational geometry without audited support context"),
], [1.35, 3.45, 2.6])
doc.add_heading("4.3 Product requirements", level=2)
add_requirement_table(doc, [
    ("GIS-PRD-01", "Authority can import KML, KMZ or zipped Shapefile into a selected tenant project and asset type.", "Successful validated import produces asset records and a published layer.", "Must"),
    ("GIS-PRD-02", "User previews geometry, feature count, fields and sample mappings before publish.", "No publish action before parse/validation succeeds.", "Must"),
    ("GIS-PRD-03", "Repeat import supports deterministic upsert and explicit replacement version.", "Counts show created/updated; previous version can be restored.", "Must"),
    ("GIS-PRD-04", "Map displays active network, asset geometry, project geofences and defects together.", "Selectable defect shows underlying asset/layer context.", "Must"),
    ("GIS-PRD-05", "Field defect and ATR evidence retains device location, accuracy, timestamp and media linkage.", "Maker/Checker journey is auditable end to end.", "Must"),
    ("GIS-PRD-06", "Responsive interaction works on 390 px mobile and desktop layouts.", "No clipped primary action; import and field actions remain operable.", "Must"),
    ("GIS-PRD-07", "Import is tenant-scoped, role-controlled and recorded in Activity Log.", "Cross-tenant reads/writes fail; publish/rollback have actor and timestamp.", "Must"),
    ("GIS-PRD-08", "mGIS/private Mappls layers can be added later without exposing OAuth secrets to browsers.", "Credential adapter boundary documented; no client secret in UI bundle.", "Should"),
])
doc.add_heading("4.4 Success measures", level=2)
add_bullet(doc, "≥95% of valid client sample features import without manual geometry conversion.")
add_bullet(doc, "100% of imported assets retain tenant, project, asset type, source identity and active layer linkage.")
add_bullet(doc, "100% of publish and rollback actions appear in the activity trail.")
add_bullet(doc, "No cross-tenant geometry, asset, defect or import-history exposure in API tests.")
add_bullet(doc, "A trained Authority User completes a 500-feature import in under five minutes, excluding file preparation.")

doc.add_heading("5. Functional Requirements (FRS)", level=1)
doc.add_heading("5.1 Import and validation", level=2)
add_requirement_table(doc, [
    ("GIS-FR-001", "File chooser accepts only .kml, .kmz and .zip; size limit is checked before parsing.", "Invalid extension/size returns a clear non-destructive message.", "Must"),
    ("GIS-FR-002", "KMZ is decompressed and the first KML document parsed. Shapefile ZIP is converted to GeoJSON using projection metadata.", "FeatureCollection is normalised to supported geometry and scalar attributes.", "Must"),
    ("GIS-FR-003", "UI lists geometry types, source fields, valid count and validation notes.", "Summary and first five mapped features are visible.", "Must"),
    ("GIS-FR-004", "Authority selects project, tenant asset type, layer name, source ID, display-name field, colour and optional replaced layer.", "Required fields gate publish; replacement list is project-scoped.", "Must"),
    ("GIS-FR-005", "Chosen source ID must be unique within the file.", "Duplicate count is shown and publish remains disabled.", "Must"),
    ("GIS-FR-006", "Server revalidates file-derived GeoJSON, project, asset type, replacement layer and maximum feature count.", "Tampered/cross-tenant requests return 4xx with no mutation.", "Must"),
])
doc.add_heading("5.2 Asset and layer publishing", level=2)
add_requirement_table(doc, [
    ("GIS-FR-010", "Each feature is matched by tenant + project + asset type + source ID.", "Existing match updates; otherwise a new asset is created.", "Must"),
    ("GIS-FR-011", "Asset name, geometry, coordinate label and up to 50 source attributes are stored.", "Register card and map reflect published values.", "Must"),
    ("GIS-FR-012", "A published layer stores source format, style, geometry type, features, version and import linkage.", "Active layer is rendered in Asset and Map & GIS screens.", "Must"),
    ("GIS-FR-013", "Replacement hides the prior active layer and increments its version number.", "Only active visible version is drawn; historical row remains stored.", "Must"),
    ("GIS-FR-014", "Import history shows source file, format/date, project/type, counts, status and rollback action.", "Authority can distinguish created versus updated assets.", "Must"),
    ("GIS-FR-015", "Rollback removes assets created by the import, restores snapshots of updated assets, hides its layer and reactivates the prior layer.", "Post-rollback register and map equal pre-import state for affected records.", "Must"),
])
doc.add_heading("5.3 Defect location and field lifecycle", level=2)
add_requirement_table(doc, [
    ("GIS-FR-020", "Defect capture records latitude, longitude, optional accuracy and geo-tagged photo/video evidence.", "Evidence retrieval remains tenant-scoped.", "Must"),
    ("GIS-FR-021", "Server computes nearest applicable asset/project geofence; client cannot claim an in-fence result.", "Distance, radius, source type and source ID are stored.", "Must"),
    ("GIS-FR-022", "Checker validates Citizen defect before Maker assignment.", "Maker cannot self-validate; invalid issue can be rejected.", "Must"),
    ("GIS-FR-023", "Maker starts assigned work and submits ATR summary, media and current device coordinates.", "ATR Submitted state is unavailable without required evidence.", "Must"),
    ("GIS-FR-024", "Checker verifies or rejects ATR; Citizen can close or reopen after resolution.", "Status, note, actor and time are retained.", "Must"),
    ("GIS-FR-025", "Map selection identifies the defect, its coordinates, geofence result, underlying asset and active network layer.", "Selected map/detail state remains consistent.", "Must"),
])

doc.add_heading("6. Software Requirements (SRS)", level=1)
doc.add_heading("6.1 Logical architecture", level=2)
add_table(doc, ["Layer", "Responsibility", "Implementation / boundary"], [
    ("Responsive React client", "File selection, parsing, preview, mapping, map display and workflow UI", "@tmcw/togeojson, fflate, shpjs, Mappls Web JS SDK adapter"),
    ("IIMM API", "Authentication, role/tenant validation, schema validation, upsert, versioning, rollback and audit", "Express + Zod; never trusts client geofence or tenancy claims"),
    ("Operational store", "Assets, layers, imports/snapshots, defects, evidence metadata and activity", "Prototype JSON store; production target must be transactional Postgres/PostGIS/object storage"),
    ("Mappls services", "Vector basemap and reverse geocoding; optional mGIS/private layers later", "Static key for current Web/REST endpoints; server-only OAuth for mGIS"),
], [1.3, 2.8, 3.3])
doc.add_heading("6.2 Data model additions", level=2)
add_table(doc, ["Entity", "Key fields", "Rules"], [
    ("GisImport", "tenantId, projectId, layerId, file/format, mapping fields, counts, status, actor/time", "Stores created IDs and pre-update asset snapshots internally for rollback; snapshots are not returned by list API."),
    ("GisLayer", "importId, supersedesLayerId, source, version, visible, style, FeatureCollection", "Exactly one intended active version per replacement chain; tenant/project validated."),
    ("Asset", "sourceId, sourceImportId, layerId, geometry, attributes", "Upsert key includes tenant, project and asset type; geometry stored as EPSG:4326 GeoJSON."),
    ("Defect", "lat/lng, accuracy, geofence, assetId, projectId, media, ATR", "Original field coordinate remains evidence; asset/network linkage is contextual."),
], [1.2, 3.15, 3.05])
doc.add_heading("6.3 API contract", level=2)
add_table(doc, ["Method / path", "Role", "Purpose", "Key response"], [
    ("GET /api/gis/imports", "Authenticated tenant role", "Tenant-scoped import history without rollback snapshots", "Safe import summaries"),
    ("POST /api/gis/imports", "Authority", "Validate and atomically publish assets + layer", "Layer and created/updated counts"),
    ("POST /api/gis/imports/:id/rollback", "Authority", "Restore pre-import asset/layer state", "Rolled-back summary"),
    ("GET /api/gis/overview", "Authenticated", "Return tenant-visible layers, assets, defects and projects", "Map-ready operational view"),
    ("GET /api/mappls/config", "Authenticated", "Provide browser static key/configuration", "Configured flag, static key, capability list"),
    ("GET /api/mappls/reverse-geocode", "Authenticated", "Resolve field coordinate to address", "Address and provider source"),
], [1.85, 1.1, 3.0, 1.45], 7.8)
doc.add_heading("6.4 Geometry and CRS rules", level=2)
add_bullet(doc, "Persist and exchange GeoJSON longitude/latitude in WGS84 (EPSG:4326).")
add_bullet(doc, "Reject empty geometry and unsupported GeometryCollection; require minimum coordinate counts for lines/polygons.")
add_bullet(doc, "Shapefile converter may reproject from its .prj definition; missing/ambiguous CRS must fail or require explicit client confirmation in a future enhancement.")
add_bullet(doc, "Calculate a display centre from feature coordinates for a human-readable asset location; this is not a survey centroid and must not replace source geometry.")

doc.add_heading("7. Security, privacy and audit", level=1)
add_requirement_table(doc, [
    ("GIS-SEC-01", "Only active Authority Users can publish or roll back GIS imports.", "API role middleware returns 403 for Maker/Checker/Citizen.", "Must"),
    ("GIS-SEC-02", "Every project, layer, asset, import and evidence lookup is tenant-scoped server-side.", "Cross-tenant test cases return 404/403 and no data.", "Must"),
    ("GIS-SEC-03", "File names and mapped properties are treated as untrusted data; no source script/HTML is executed.", "KML is parsed as XML/GeoJSON and rendered as text/geometry only.", "Must"),
    ("GIS-SEC-04", "Client receives only a Mappls browser static key; OAuth client secret and mGIS token generation remain server-side.", "No client-secret string in bundle, response or logs.", "Must"),
    ("GIS-SEC-05", "Publish and rollback create activity entries with actor, role, entity, timestamp and detail.", "Activity Log contains IMPORTED_GIS_ASSETS / ROLLED_BACK_GIS_IMPORT.", "Must"),
    ("GIS-SEC-06", "Rollback snapshots are internal and excluded from import-history responses.", "GET response has no assetSnapshots or createdAssetIds fields.", "Must"),
    ("GIS-SEC-07", "Production storage encrypts sensitive evidence at rest and uses signed/authorised media delivery.", "Architecture/security review before production.", "Must"),
])
doc.add_heading("7.1 Mappls credential decision", level=2)
add_note(doc, "Current integration", "A Mappls Console static key is correct for the Web Map JS SDK and current reverse-geocoding endpoint when those products are enabled for the same Console application. The deployment accepts MAPPLS_ACCESS_TOKEN and the lowercase mappls_access_token alias. Whitelist the deployed domain for browser use. If private mGIS methods are adopted, add server-held OAuth client credentials/token; never send the client secret to the browser.")
doc.add_heading("7.2 Data retention questions", level=2)
add_bullet(doc, "How long must source file, parsed features, import snapshots and rolled-back history be retained?")
add_bullet(doc, "Are citizen/field coordinates personal data under the client’s policy, and what masking applies in reports?")
add_bullet(doc, "Must uploaded media carry a visible watermark, hash, capture timestamp and tamper-evidence record?")
doc.add_page_break()

doc.add_heading("8. Non-functional requirements", level=1)
add_requirement_table(doc, [
    ("GIS-NFR-01", "Import 5,000 supported features without browser crash on agreed minimum device.", "Performance test on client device/browser matrix.", "Must"),
    ("GIS-NFR-02", "API publish is atomic: no half-created asset set or visible layer on validation failure.", "Fault injection and store transaction test in production architecture.", "Must"),
    ("GIS-NFR-03", "Map and import UI remain usable at 390 px width and current desktop browsers.", "Responsive visual QA with no clipped primary controls.", "Must"),
    ("GIS-NFR-04", "Import feedback identifies file, format, feature count, warnings and exact blocking issue.", "User can correct mapping without reloading application.", "Must"),
    ("GIS-NFR-05", "Operational APIs return within 2 s p95 excluding Mappls and large import processing; map initial content within 4 s p95.", "Production observability dashboard and load test.", "Should"),
    ("GIS-NFR-06", "Offline field edits follow existing server-timestamp-wins rule; conflicts enter manual review.", "No silent drop of conflicting local change.", "Must"),
    ("GIS-NFR-07", "Map remains operational with a geometry fallback if Mappls SDK/static key is unavailable.", "Network/defect context still renders with a configuration notice.", "Must"),
    ("GIS-NFR-08", "Production target supports transactional PostGIS and object storage; JSON store is prototype-only.", "Architecture readiness gate before go-live.", "Must"),
])
doc.add_heading("8.1 Accessibility and usability", level=2)
add_bullet(doc, "All primary actions have text labels; file chooser, fields, tables and modal are keyboard-addressable.")
add_bullet(doc, "Status is not conveyed by colour alone; badges and validation messages include text.")
add_bullet(doc, "Map information required to complete a workflow is duplicated in structured asset/defect details.")
add_bullet(doc, "Target WCAG 2.1 AA contrast and form labelling for the production release.")
doc.add_page_break()

doc.add_heading("9. Acceptance and traceability", level=1)
add_table(doc, ["Scenario", "Precondition / action", "Expected result", "Req. IDs"], [
    ("A1 KML publish", "Authority uploads valid WGS84 KML, selects project/type, maps asset_id and name, publishes.", "Preview shows geometry; N assets and one active layer created; activity logged.", "PRD-01/02, FR-001–012"),
    ("A2 Repeat version", "Authority imports revised file and selects prior layer as replacement.", "Matching source IDs update, new IDs create, layer version increments and prior hides.", "PRD-03, FR-010–014"),
    ("A3 Rollback", "Authority rolls back A2.", "A2-created assets removed, updated assets restored, A2 layer hidden, prior layer visible.", "FR-015, SEC-05/06"),
    ("A4 Duplicate IDs", "File contains duplicate values in selected source-ID field.", "Publish disabled; duplicate count shown; no server mutation.", "FR-005/006"),
    ("A5 Tenant isolation", "NHAI authority attempts PWD project/replacement layer.", "Request rejected; no import/asset/layer visible across tenant.", "SEC-01/02"),
    ("A6 Defect lifecycle", "Citizen reports with coordinates/media; Checker validates; Maker submits geo-ATR; Checker verifies.", "Every actor/state/evidence/geofence retained and map shows underlying asset/network.", "FR-020–025"),
    ("A7 Mobile layout", "Authority opens Assets/import at 390×844; Maker opens field actions.", "Actions, modal and data remain usable without horizontal page overflow.", "PRD-06, NFR-03"),
    ("A8 Credential fallback", "Mappls key absent or SDK fails.", "Fallback geometry renders; clear configuration message shown; rest of workflow works.", "NFR-07"),
], [1.05, 2.7, 2.8, 1.05], 7.5)
doc.add_heading("9.1 Current implementation evidence", level=2)
add_bullet(doc, "Automated TypeScript, API, workflow and production-build checks pass on the feature branch.")
add_bullet(doc, "API test covers create/update versioning, safe import-history response and rollback restoration.")
add_bullet(doc, "Browser QA successfully parsed and published the included two-feature KML sample.")
add_bullet(doc, "Responsive QA completed at 390×844 for Asset Management and the import modal.")
add_bullet(doc, "Implementation remains unmerged and can be revised or discarded without altering the production baseline.")
doc.add_page_break()

doc.add_heading("10. Client decisions required", level=1)
add_table(doc, ["Decision", "Question for client", "Proposed default", "Impact if changed"], [
    ("DEC-01", "Which source formats and archive contents are contractually supported?", "KML/KMZ/SHP ZIP", "Parser, validation and support scope"),
    ("DEC-02", "Which CRS values are accepted for Shapefile, and what happens when .prj is missing?", "Require valid .prj; output EPSG:4326", "Import success and coordinate liability"),
    ("DEC-03", "What uniquely identifies an asset across files and versions?", "Configured source-ID field per import", "Duplicate/upsert correctness"),
    ("DEC-04", "Is one feature always one asset, or can a layer be reference-only?", "One feature → one asset + layer feature", "Register volume and workflow linkage"),
    ("DEC-05", "Must a new file fully replace a prior layer, or may versions coexist?", "One active replacement version; history retained", "Map ambiguity and rollback model"),
    ("DEC-06", "Should invalid features block the whole import or be quarantined and partially published?", "Block atomic publish after preview", "Operational safety versus throughput"),
    ("DEC-07", "What are the defect asset-buffer, project-radius and network-snap tolerances?", "75 m / project setting / 25 m", "Assignment suggestions and exception volume"),
    ("DEC-08", "Does outside-geofence defect capture block, warn, or require supervisor override?", "Warn + audit; do not block", "Citizen/field completeness and fraud control"),
    ("DEC-09", "What GPS accuracy and media evidence are mandatory for defect and ATR?", "Warn above 100 m; photo/video required by workflow", "Field usability and evidentiary strength"),
    ("DEC-10", "Will the client supply Mappls static-key domain/IP restrictions and any mGIS OAuth app?", "Static key now; mGIS later", "Basemap/reverse-geocode/private layers"),
    ("DEC-11", "What import size, feature volume, processing SLA and concurrent-user targets apply?", "10/25 MB, 5,000 features", "Architecture and cost"),
    ("DEC-12", "Retention, deletion and audit duration for source files, snapshots, locations and media?", "Client policy required", "Privacy, storage and rollback window"),
], [0.65, 3.0, 2.0, 1.6], 7.4)
doc.add_heading("10.1 Review response template", level=2)
doc.add_paragraph("For each decision, record: Approved default / Revised value / Deferred / Out of scope, plus owner and due date. Unresolved Must-level decisions block merge to production.")
add_table(doc, ["Decision ID", "Client response", "Owner", "Due date"], [("", "", "", "") for _ in range(6)], [1.0, 3.8, 1.3, 1.1])
doc.add_page_break()

doc.add_heading("11. Rollback and change control", level=1)
doc.add_heading("11.1 Code isolation", level=2)
doc.add_paragraph("All new GIS file-import code, tests, fixture and this specification are maintained on branch feature/kml-shp-gis-import. The branch is pushed independently and is not merged to main. The currently published Replit deployment follows main at commit 2fa44ad.")
add_table(doc, ["Client outcome", "Delivery action", "Production effect"], [
    ("Approve", "Update document to v1.0, complete UAT/security gates, merge reviewed branch through PR.", "Feature becomes part of next approved release."),
    ("Approve with changes", "Revise requirements and code on same branch; rerun traceability and tests; request review again.", "No effect until later approved merge."),
    ("Reject / defer", "Close or retain branch for reference; do not merge.", "Production remains on current main baseline."),
    ("Post-merge incident", "Revert the merge commit or disable feature via release control; use per-import rollback for data-level issue.", "Code and data rollback paths remain separate."),
], [1.5, 3.75, 2.2])
doc.add_heading("11.2 Data rollback", level=2)
for step_number, step_text in enumerate([
    "Authority selects the published import history record and confirms rollback.",
    "Server removes assets created only by that import and restores snapshots for assets it updated.",
    "Imported layer is hidden; its superseded layer is made visible again when present.",
    "Import status becomes Rolled back and the action is written to Activity Log.",
    "Subsequent operational records linked to imported assets must be reviewed before production rollback; v0.1 does not cascade-delete defects or inspections.",
], 1):
    doc.add_paragraph(f"{step_number}.  {step_text}")
add_note(doc, "Production caution", "The prototype rollback uses stored asset snapshots. A production PostGIS design must execute import and rollback in database transactions, protect referential integrity, and define how defects/inspections linked after import are handled.", "FFF3D6")
doc.add_heading("11.3 Merge readiness gates", level=2)
for text in [
    "Client resolves DEC-01 through DEC-12 and approves requirements baseline.",
    "Client sample KML/KMZ/SHP files pass CRS/attribute/volume UAT.",
    "Security review approves credential handling, file parsing, tenancy and media retention.",
    "Production persistence migrates from JSON store to transactional database/object storage.",
    "Performance, accessibility and mobile device matrix meet agreed NFRs.",
    "A pull request is reviewed; no direct merge from the experimental branch.",
]: add_bullet(doc, text)

doc.add_page_break()
doc.add_heading("Appendix A — Mappls credential matrix", level=1)
add_table(doc, ["Capability", "Credential proposed", "Where used", "Configuration check"], [
    ("Mappls Web Map JS SDK v3", "Console static key / access_token", "Browser", "Enable Web Maps; whitelist mappls-dic-iimm.replit.app and approved domains."),
    ("Current Reverse Geocoding REST API", "Same Console static key when API enabled", "IIMM server", "Enable Reverse Geocoding; verify server access restrictions suitable for Replit."),
    ("Mappls mGIS methods/private layers", "OAuth2 access token generated from server-held client credentials", "Server adapter only", "Separate mGIS application/credentials; token typically short-lived; never expose secret."),
    ("Other premium/reserved REST APIs", "Per current endpoint documentation / Console enablement", "Prefer server", "Confirm product entitlement and endpoint-specific auth before implementation."),
], [1.9, 1.9, 1.2, 2.5], 8)
doc.add_heading("Appendix B — Implemented branch inventory", level=1)
add_table(doc, ["Area", "Implementation"], [
    ("Parser", "Browser-side KML, KMZ and zipped Shapefile conversion to normalised GeoJSON."),
    ("Authority UX", "Import modal, validation summary, field mapping, Mappls preview, sample rows and publish."),
    ("Server", "Tenant-scoped schema validation, asset upsert, layer versioning, safe history and rollback."),
    ("Operations map", "Imported active layer and assets appear below selectable defects with underlying network context."),
    ("Quality", "Automated API lifecycle test, two-feature KML fixture, desktop/mobile browser QA and production build."),
], [1.4, 5.9], 8.5)
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = end.add_run("END OF CLIENT REVIEW DRAFT")
rr.bold = True
rr.font.size = Pt(9)
rr.font.color.rgb = RGBColor.from_string(NAVY)
add_rule(end, GOLD, 18)

doc.core_properties.title = "IIMM GIS Asset & Network Import — PRD/FRS/SRS v0.1"
doc.core_properties.subject = "Client review draft for KML/KMZ/Shapefile import and geotagged defect workflows"
doc.core_properties.author = "IIMM Delivery Team"
doc.core_properties.keywords = "IIMM, Mappls, GIS, KML, KMZ, Shapefile, PRD, FRS, SRS"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
